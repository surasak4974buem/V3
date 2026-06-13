# -*- coding: utf-8 -*-
"""
สร้างคู่มือการใช้งาน ระบบบริหารคลังอะไหล่ รพ.นครพิงค์ V3
ออกไฟล์: NKP_Manual_V3.docx และ NKP_Manual_V3.pdf
"""

import os, sys
from pathlib import Path
from datetime import date

BASE    = Path(r"D:\Cowork\NKPmedparts2026V2")
IMG_DIR = BASE / "รูปสำหรับทำคู่มือ"
OUT_DIR = BASE

# ─── Image mapping ───────────────────────────────────────────────
IMGS = {
    "login":        IMG_DIR / "1.png",
    "login2":       IMG_DIR / "2.png",
    "dashboard":    IMG_DIR / "3.png",
    "catalog":      IMG_DIR / "4.png",
    "transaction":  IMG_DIR / "5.png",
    "ledger":       IMG_DIR / "6.png",
    "planning":     IMG_DIR / "7.png",
    "procurement":  IMG_DIR / "8.png",
    "settings":     IMG_DIR / "9.png",
    "usermgmt":     IMG_DIR / "10.png",
    "budget":       IMG_DIR / "11.png",
    "genuser1":     IMG_DIR / "General User01.png",
    "genuser2":     IMG_DIR / "General User02.png",
    "stockmgr":     IMG_DIR / "Stock Manager.png",
}

TODAY_TH = "13 มิถุนายน 2568"
VERSION  = "3.0"
DOC_CODE = "SMM-NKP-07-1-2568"

# ════════════════════════════════════════════════════════════════
#  PART 1 – PDF  (ReportLab)
# ════════════════════════════════════════════════════════════════
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm, cm
from reportlab.lib import colors
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Image as RLImage,
                                 Table, TableStyle, PageBreak, HRFlowable,
                                 KeepTogether, BaseDocTemplate, PageTemplate, Frame,
                                 NextPageTemplate)
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.graphics.shapes import Drawing, Rect, String
from reportlab.graphics import renderPDF

# ── Register Thai fonts ─────────────────────────────────────────
FONT_DIR = Path(r"C:\Windows\Fonts")
pdfmetrics.registerFont(TTFont("THSarabun",     FONT_DIR / "cordia.ttc",  subfontIndex=0))
pdfmetrics.registerFont(TTFont("THSarabunBold", FONT_DIR / "cordia.ttc",  subfontIndex=1))
pdfmetrics.registerFont(TTFont("Angsana",       FONT_DIR / "angsana.ttc", subfontIndex=0))
pdfmetrics.registerFont(TTFont("AngsanaBold",   FONT_DIR / "angsana.ttc", subfontIndex=1))

# ── Color palette ───────────────────────────────────────────────
C_NAVY   = colors.HexColor("#003580")
C_BLUE   = colors.HexColor("#0052A3")
C_LBLUE  = colors.HexColor("#006EC0")
C_GOLD   = colors.HexColor("#C8A415")
C_GREEN  = colors.HexColor("#1A7A4A")
C_RED    = colors.HexColor("#C0392B")
C_LGRAY  = colors.HexColor("#F5F7FA")
C_BORDER = colors.HexColor("#DCE3ED")
C_MUTED  = colors.HexColor("#6B7A8D")
C_WHITE  = colors.white
C_BLACK  = colors.black
C_TEAL   = colors.HexColor("#0099CC")

W, H = A4  # 595.3 x 841.9 pt

# ── Styles ──────────────────────────────────────────────────────
def S(name, **kw):
    defaults = dict(fontName="THSarabun", fontSize=14, leading=20, textColor=C_BLACK)
    defaults.update(kw)
    return ParagraphStyle(name, **defaults)

sBody   = S("Body",   fontSize=13, leading=19, alignment=TA_JUSTIFY)
sSmall  = S("Small",  fontSize=11, leading=16, textColor=C_MUTED)
sCapt   = S("Caption",fontSize=11, leading=15, textColor=C_MUTED, alignment=TA_CENTER, fontName="THSarabun")
sH1     = S("H1",     fontSize=20, leading=26, textColor=C_BLUE,  fontName="THSarabunBold", spaceAfter=6)
sH2     = S("H2",     fontSize=16, leading=22, textColor=C_NAVY,  fontName="THSarabunBold", spaceAfter=4)
sH3     = S("H3",     fontSize=14, leading=20, textColor=C_BLUE,  fontName="THSarabunBold")
sBullet = S("Bullet", fontSize=13, leading=19, leftIndent=16, firstLineIndent=0,
            bulletIndent=0, spaceAfter=3)
sTH     = S("TH",     fontSize=12, leading=16, fontName="THSarabunBold", textColor=C_WHITE, alignment=TA_CENTER)
sTD     = S("TD",     fontSize=12, leading=16, fontName="THSarabun")
sNum    = S("Num",    fontSize=13, leading=19, leftIndent=20)
sCover  = S("Cover",  fontSize=13, textColor=C_WHITE, alignment=TA_CENTER)
sCenter = S("Center", fontSize=13, leading=19, alignment=TA_CENTER)

def bold(txt, color=None):
    c = f' color="{color}"' if color else ""
    return f'<font name="THSarabunBold"{c}>{txt}</font>'

def img_flow(key, caption="", width_mm=155, max_h_mm=100):
    from PIL import Image as PILImage
    p = IMGS.get(key)
    if not p or not p.exists():
        return []
    with PILImage.open(str(p)) as im:
        iw, ih = im.size
    ratio = ih / iw
    w = width_mm * mm
    h = w * ratio
    if h > max_h_mm * mm:
        h = max_h_mm * mm
        w = h / ratio
    ri = RLImage(str(p), width=w, height=h)
    ri.hAlign = "CENTER"
    elems = [Spacer(1, 2*mm), ri]
    if caption:
        elems.append(Paragraph(f"ภาพที่: {caption}", sCapt))
    elems.append(Spacer(1, 4*mm))
    return elems

def hr():
    return HRFlowable(width="100%", thickness=1, color=C_BORDER, spaceAfter=6, spaceBefore=6)

def info_box(title, body, color=C_TEAL):
    data = [[Paragraph(f'<font name="THSarabunBold" color="{color.hexval()}">{title}</font>', S("ib")),
             Paragraph(body, sBody)]]
    t = Table(data, colWidths=[40*mm, 120*mm])
    t.setStyle(TableStyle([
        ("VALIGN",    (0,0), (-1,-1), "TOP"),
        ("LEFTPADDING",(0,0),(-1,-1), 10),
        ("RIGHTPADDING",(0,0),(-1,-1), 10),
        ("TOPPADDING", (0,0),(-1,-1), 8),
        ("BOTTOMPADDING",(0,0),(-1,-1), 8),
        ("BACKGROUND",(0,0),(-1,-1), colors.HexColor("#EEF6FF")),
        ("LINEBEFORE",(0,0),(0,-1), 4, color),
        ("ROUNDEDCORNERS", [4]),
    ]))
    return t

def step_table(steps):
    rows = []
    for i, (title, desc) in enumerate(steps, 1):
        rows.append([
            Paragraph(str(i), ParagraphStyle("SN", fontName="THSarabunBold", fontSize=13,
                                              textColor=C_WHITE, alignment=TA_CENTER)),
            Paragraph(f'{bold(title)}<br/>{desc}', sBody)
        ])
    t = Table(rows, colWidths=[12*mm, 148*mm])
    style = [("VALIGN",(0,0),(-1,-1),"TOP"),
             ("TOPPADDING",(0,0),(-1,-1),6),
             ("BOTTOMPADDING",(0,0),(-1,-1),6),
             ("LEFTPADDING",(1,0),(1,-1),10),
             ("BACKGROUND",(0,0),(0,-1),C_BLUE),
             ("ROWBACKGROUNDS",(1,0),(1,-1),[C_WHITE, C_LGRAY]),
             ("ROUNDEDCORNERS",[4]),
             ("GRID",(0,0),(-1,-1),0.5,C_BORDER)]
    t.setStyle(TableStyle(style))
    return t

def role_chip(role):
    colors_map = {"Admin":"#C8A415","Stock Manager":"#0052A3","General User":"#6B7A8D"}
    c = colors_map.get(role, "#6B7A8D")
    return f'<font name="THSarabunBold" color="{c}">● {role}</font>'

# ─── Cover page builder ─────────────────────────────────────────
def build_cover_page(canvas, doc):
    canvas.saveState()
    # Gradient background (simulate with rectangles)
    steps = 30
    for i in range(steps):
        r = (0 + i*0.01)
        g = (0.21 + i*0.012)
        b = (0.50 + i*0.015)
        canvas.setFillColorRGB(min(r,1), min(g,1), min(b,1))
        canvas.rect(0, H - H*(i+1)/steps, W, H/steps + 1, stroke=0, fill=1)

    # Gold decorative strip
    canvas.setFillColor(C_GOLD)
    canvas.rect(0, H*0.62, W, 3, stroke=0, fill=1)

    # Hospital emblem circle
    canvas.setFillColor(colors.HexColor("#FFFFFF20"))
    canvas.circle(W*0.75, H*0.75, 120, stroke=0, fill=1)
    canvas.setFillColor(colors.HexColor("#FFFFFF10"))
    canvas.circle(W*0.25, H*0.3,  80,  stroke=0, fill=1)

    # Logo area (cross symbol)
    cx, cy = W/2, H*0.78
    canvas.setFillColor(C_WHITE)
    canvas.setFont("THSarabunBold", 52)
    canvas.drawCentredString(cx, cy, "✚")

    # Hospital name
    canvas.setFont("THSarabunBold", 20)
    canvas.setFillColor(colors.HexColor("#7ECEF4"))
    canvas.drawCentredString(cx, H*0.73, "โรงพยาบาลนครพิงค์  Nakornping Hospital")

    canvas.setFont("THSarabun", 14)
    canvas.setFillColor(colors.HexColor("#AED6F1"))
    canvas.drawCentredString(cx, H*0.70, "ศูนย์เครื่องมือแพทย์  •  งานซ่อมบำรุงและคลังอะไหล่")

    # Gold line
    canvas.setStrokeColor(C_GOLD)
    canvas.setLineWidth(2)
    canvas.line(cx-80, H*0.675, cx+80, H*0.675)

    # Main title
    canvas.setFont("THSarabunBold", 34)
    canvas.setFillColor(C_WHITE)
    canvas.drawCentredString(cx, H*0.63, "คู่มือการใช้งาน")
    canvas.setFont("THSarabunBold", 26)
    canvas.drawCentredString(cx, H*0.595, "ระบบบริหารคลังอะไหล่และอุปกรณ์")
    canvas.setFont("THSarabun", 17)
    canvas.setFillColor(colors.HexColor("#90CAF9"))
    canvas.drawCentredString(cx, H*0.565, "Medical Spare Parts Management System")

    # Badges
    def badge(text, y):
        tw = canvas.stringWidth(text, "THSarabun", 13) + 30
        canvas.setFillColor(colors.HexColor("#FFFFFF18"))
        canvas.roundRect(cx - tw/2, y-4, tw, 22, 11, stroke=0, fill=1)
        canvas.setFont("THSarabun", 13)
        canvas.setFillColor(colors.HexColor("#D4E8FF"))
        canvas.drawCentredString(cx, y+4, text)

    badge(f"⭐  Version {VERSION}   |   รหัสเอกสาร: {DOC_CODE}", H*0.52)
    badge(f"📅  วันที่จัดทำ: {TODAY_TH}   |   มาตรฐาน SMM 07-1:2024", H*0.49)

    # Bottom bar
    canvas.setFillColor(colors.HexColor("#FFFFFF15"))
    canvas.rect(0, 0, W, 50, stroke=0, fill=1)
    canvas.setFont("THSarabun", 11)
    canvas.setFillColor(colors.HexColor("#90CAF9"))
    canvas.drawCentredString(cx, 20, "จัดทำโดย: กลุ่มงานวิศวกรรมการแพทย์  |  โรงพยาบาลนครพิงค์  จ.เชียงใหม่")

    canvas.restoreState()

def build_back_cover(canvas, doc):
    canvas.saveState()
    # Dark navy background
    canvas.setFillColor(C_NAVY)
    canvas.rect(0, 0, W, H, stroke=0, fill=1)

    # Gold strip top
    canvas.setFillColor(C_GOLD)
    canvas.rect(0, H-8, W, 8, stroke=0, fill=1)
    canvas.rect(0, 0,   W, 8, stroke=0, fill=1)

    # Center content
    cx = W/2
    canvas.setFont("THSarabunBold", 28)
    canvas.setFillColor(C_WHITE)
    canvas.drawCentredString(cx, H*0.65, "ระบบบริหารคลังอะไหล่")
    canvas.setFont("THSarabun", 16)
    canvas.setFillColor(colors.HexColor("#90CAF9"))
    canvas.drawCentredString(cx, H*0.60, "Medical Spare Parts Management System")

    canvas.setStrokeColor(C_GOLD)
    canvas.setLineWidth(1.5)
    canvas.line(cx-120, H*0.575, cx+120, H*0.575)

    canvas.setFont("THSarabun", 13)
    canvas.setFillColor(colors.HexColor("#AED6F1"))
    for i, line in enumerate([
        "โรงพยาบาลนครพิงค์  จังหวัดเชียงใหม่",
        "ศูนย์เครื่องมือแพทย์  กลุ่มงานวิศวกรรมการแพทย์",
        "",
        f"รหัสเอกสาร: {DOC_CODE}",
        f"Version {VERSION}  |  {TODAY_TH}",
        "",
        "GitHub: surasak4974buem/V3",
    ]):
        canvas.drawCentredString(cx, H*0.54 - i*18, line)

    # QR placeholder box
    canvas.setFillColor(colors.HexColor("#FFFFFF10"))
    canvas.roundRect(cx-35, H*0.2, 70, 70, 8, stroke=0, fill=1)
    canvas.setFont("THSarabun", 10)
    canvas.setFillColor(colors.HexColor("#6B7A8D"))
    canvas.drawCentredString(cx, H*0.205, "เข้าใช้ระบบออนไลน์")

    canvas.restoreState()

# ─── Page numbering ─────────────────────────────────────────────
def normal_page(canvas, doc):
    canvas.saveState()
    # Header
    canvas.setFillColor(C_BLUE)
    canvas.rect(0, H-14*mm, W, 14*mm, stroke=0, fill=1)
    canvas.setFont("THSarabun", 10)
    canvas.setFillColor(C_WHITE)
    canvas.drawString(15*mm, H-9*mm, f"คู่มือการใช้งาน  ระบบบริหารคลังอะไหล่  รพ.นครพิงค์  V{VERSION}")
    canvas.drawRightString(W-15*mm, H-9*mm, f"รหัส: {DOC_CODE}")
    # Footer
    canvas.setFillColor(C_LGRAY)
    canvas.rect(0, 0, W, 10*mm, stroke=0, fill=1)
    canvas.setFont("THSarabun", 10)
    canvas.setFillColor(C_MUTED)
    canvas.drawString(15*mm, 3*mm, "กลุ่มงานวิศวกรรมการแพทย์  โรงพยาบาลนครพิงค์  จ.เชียงใหม่")
    canvas.drawRightString(W-15*mm, 3*mm, f"หน้าที่ {doc.page}")
    canvas.restoreState()

# ─── Build PDF ──────────────────────────────────────────────────
def build_pdf():
    out = OUT_DIR / "NKP_Manual_V3.pdf"
    doc = BaseDocTemplate(
        str(out), pagesize=A4,
        leftMargin=20*mm, rightMargin=20*mm,
        topMargin=20*mm, bottomMargin=16*mm,
        title=f"คู่มือการใช้งาน ระบบบริหารคลังอะไหล่ รพ.นครพิงค์ V{VERSION}",
        author="กลุ่มงานวิศวกรรมการแพทย์ โรงพยาบาลนครพิงค์",
        subject="User Manual - Medical Spare Parts Management System",
    )

    cover_frame  = Frame(0, 0, W, H, leftPadding=0, rightPadding=0, topPadding=0, bottomPadding=0)
    normal_frame = Frame(20*mm, 16*mm, W-40*mm, H-36*mm)

    doc.addPageTemplates([
        PageTemplate(id="Cover",  frames=[cover_frame],  onPage=build_cover_page),
        PageTemplate(id="Back",   frames=[cover_frame],  onPage=build_back_cover),
        PageTemplate(id="Normal", frames=[normal_frame], onPage=normal_page),
    ])

    story = []
    def next_tpl(name):
        story.append(NextPageTemplate(name))
        story.append(PageBreak())

    def section_header(num, th, en, icon=""):
        data = [[Paragraph(f'{icon} {num}', ParagraphStyle("SN2", fontName="THSarabunBold",
                                                              fontSize=18, textColor=C_WHITE,
                                                              alignment=TA_CENTER)),
                 [Paragraph(bold(th, "#003580"), sH1),
                  Paragraph(en, S("En", fontSize=12, textColor=C_MUTED))]]]
        t = Table(data, colWidths=[20*mm, 140*mm])
        t.setStyle(TableStyle([
            ("BACKGROUND",(0,0),(0,0), C_BLUE),
            ("VALIGN",(0,0),(-1,-1),"MIDDLE"),
            ("LEFTPADDING",(0,0),(-1,-1),8),
            ("TOPPADDING",(0,0),(-1,-1),8),
            ("BOTTOMPADDING",(0,0),(-1,-1),8),
            ("LINEBELOW",(0,0),(-1,-1),2,C_GOLD),
        ]))
        story.append(t)
        story.append(Spacer(1, 6*mm))

    def perm_table():
        hdr = ["ฟังก์ชัน / Feature","Admin","Stock Manager","General User"]
        rows = [
            hdr,
            ["ดูแดชบอร์ด","✅","✅","✅"],
            ["เพิ่ม/แก้ไข/ลบอะไหล่","✅","✅","❌"],
            ["บันทึกรับเข้า/เบิกออก","✅","✅","✅"],
            ["แผนบำรุงรักษา PM","✅","✅","👁️ ดูได้"],
            ["ใบขอจัดซื้อ","✅","✅","❌"],
            ["แผนงบประมาณ","✅","✅","❌"],
            ["จัดการผู้ใช้","✅","❌","❌"],
            ["ตั้งค่าระบบ / LINE OA","✅","❌","❌"],
            ["Cloud Sync (Google Sheets)","✅","❌","❌"],
        ]
        col_w = [80*mm, 25*mm, 35*mm, 30*mm]
        t = Table([[Paragraph(c if i==0 else c, sTH if r_i==0 else sTD) for i,c in enumerate(row)]
                    for r_i, row in enumerate(rows)], colWidths=col_w)
        t.setStyle(TableStyle([
            ("BACKGROUND",(0,0),(-1,0), C_BLUE),
            ("TEXTCOLOR",(0,0),(-1,0), C_WHITE),
            ("ALIGN",(1,0),(-1,-1),"CENTER"),
            ("FONTNAME",(0,0),(-1,0),"THSarabunBold"),
            ("ROWBACKGROUNDS",(0,1),(-1,-1),[C_WHITE,C_LGRAY]),
            ("GRID",(0,0),(-1,-1),0.5,C_BORDER),
            ("TOPPADDING",(0,0),(-1,-1),5),
            ("BOTTOMPADDING",(0,0),(-1,-1),5),
        ]))
        return t

    # ── Cover (placeholder — actual cover drawn by onPage) ──────
    story.append(Paragraph(" ", S("sp")))   # keep cover template active

    # ── Switch to Normal ────────────────────────────────────────
    next_tpl("Normal")

    # ══ คำนำ ══════════════════════════════════════════════════════
    story.append(Paragraph(bold("คำนำ", "#003580"), sH1))
    story.append(HRFlowable(width="100%", thickness=2, color=C_GOLD, spaceAfter=8))
    story.append(Paragraph(
        "คู่มือการใช้งาน ระบบบริหารคลังอะไหล่และอุปกรณ์ โรงพยาบาลนครพิงค์ (Version 3) "
        "จัดทำขึ้นเพื่อเป็นแนวทางสำหรับเจ้าหน้าที่ วิศวกรชีวการแพทย์ และผู้ดูแลระบบ "
        "ในการใช้งานระบบสารสนเทศเพื่อการบริหารจัดการคลังอะไหล่เครื่องมือแพทย์ให้เป็นไปตาม "
        "มาตรฐานการบริหารงานซ่อมบำรุงเครื่องมือแพทย์ (SMM 07-1:2024) และข้อกำหนด "
        "ของกระทรวงสาธารณสุขและมาตรฐาน ISO 13485:2016", sBody))
    story.append(Spacer(1,4*mm))
    story.append(Paragraph(
        "ระบบนี้พัฒนาขึ้นในรูปแบบ Web Application ที่ไม่ต้องติดตั้งซอฟต์แวร์เพิ่มเติม "
        "สามารถเข้าใช้งานผ่าน Browser ได้ทุกอุปกรณ์ และเชื่อมต่อกับ Google Sheets "
        "เป็นฐานข้อมูลคลาวด์ พร้อมระบบแจ้งเตือนผ่าน LINE OA", sBody))
    story.append(Spacer(1,4*mm))
    story.append(Paragraph(bold("วัตถุประสงค์ของคู่มือ"), sH3))
    for b in ["อธิบายขั้นตอนการใช้งานระบบในแต่ละโมดูลอย่างละเอียด พร้อมภาพประกอบ",
              "กำหนดสิทธิ์การเข้าถึงและความรับผิดชอบของผู้ใช้แต่ละระดับ",
              "แนะนำการแก้ไขปัญหาเบื้องต้นและข้อควรระวังในการใช้งาน",
              "ใช้เป็นเอกสารอ้างอิงสำหรับการฝึกอบรมบุคลากรใหม่"]:
        story.append(Paragraph(f"• {b}", sBullet))
    story.append(Spacer(1,6*mm))
    story.append(Paragraph(bold("ผู้จัดทำและผู้รับผิดชอบ"), sH3))
    meta = [["รายการ","รายละเอียด"],
            ["ชื่อระบบ","ระบบบริหารคลังอะไหล่และอุปกรณ์ รพ.นครพิงค์"],
            ["รหัสเอกสาร", DOC_CODE],
            ["เวอร์ชัน", f"Version {VERSION}"],
            ["วันที่จัดทำ", TODAY_TH],
            ["ผู้จัดทำ","กลุ่มงานวิศวกรรมการแพทย์ โรงพยาบาลนครพิงค์"],
            ["มาตรฐานอ้างอิง","SMM 07-1:2024, ISO 13485:2016, กฎกระทรวงสาธารณสุข"]]
    mt = Table([[Paragraph(r[0],sTH if i==0 else sTD),
                  Paragraph(r[1],sTH if i==0 else sTD)] for i,r in enumerate(meta)],
               colWidths=[55*mm, 105*mm])
    mt.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(-1,0),C_BLUE),
        ("TEXTCOLOR",(0,0),(-1,0),C_WHITE),
        ("GRID",(0,0),(-1,-1),0.5,C_BORDER),
        ("ROWBACKGROUNDS",(0,1),(-1,-1),[C_WHITE,C_LGRAY]),
        ("TOPPADDING",(0,0),(-1,-1),5),("BOTTOMPADDING",(0,0),(-1,-1),5),
    ]))
    story.append(mt)
    story.append(PageBreak())

    # ══ สารบัญ ═══════════════════════════════════════════════════
    story.append(Paragraph(bold("สารบัญ  Table of Contents", "#003580"), sH1))
    story.append(HRFlowable(width="100%", thickness=2, color=C_GOLD, spaceAfter=8))
    toc_items = [
        ("คำนำ","Preface","I"),
        ("สารบัญ","Table of Contents","II"),
        ("1","ภาพรวมระบบ — System Overview","3"),
        ("2","การเข้าสู่ระบบ — Login & Authentication","4"),
        ("3","แดชบอร์ดภาพรวม — Dashboard","5"),
        ("4","คลังอะไหล่และอุปกรณ์ — Parts Catalog","6"),
        ("5","บันทึกรับเข้า/เบิกออก — Transactions","7"),
        ("6","ประวัติรายการคลัง — Ledger Logs","8"),
        ("7","แผนบำรุงรักษาเชิงรุก — PM Planning","9"),
        ("8","ใบขอจัดซื้ออะไหล่ — Procurement","10"),
        ("9","แผนงบประมาณประจำปี — Budget Management","11"),
        ("10","จัดการสิทธิ์ผู้ใช้ — User Management","12"),
        ("11","ตั้งค่าระบบ — Cloud & LINE Settings","13"),
        ("12","สิทธิ์การใช้งาน & แก้ปัญหา — Permissions & FAQ","14"),
    ]
    toc_data = []
    for num, title, pg in toc_items:
        is_hdr = num in ("คำนำ","สารบัญ")
        toc_data.append([
            Paragraph(num, ParagraphStyle("TN",fontName="THSarabunBold",fontSize=12,
                                           textColor=C_WHITE,alignment=TA_CENTER)),
            Paragraph(title, S("TT",fontSize=13,fontName="THSarabunBold" if is_hdr else "THSarabun")),
            Paragraph(pg, S("TP",fontSize=12,textColor=C_MUTED,alignment=TA_RIGHT)),
        ])
    tt = Table(toc_data, colWidths=[15*mm, 130*mm, 15*mm],
               rowHeights=[None]*len(toc_data))
    tt.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(0,-1),C_BLUE),
        ("ROWBACKGROUNDS",(1,0),(1,-1),[C_WHITE,C_LGRAY]),
        ("LINEBELOW",(0,0),(-1,-1),0.5,C_BORDER),
        ("TOPPADDING",(0,0),(-1,-1),5),("BOTTOMPADDING",(0,0),(-1,-1),5),
        ("LEFTPADDING",(1,0),(1,-1),10),
    ]))
    story.append(tt)
    story.append(PageBreak())

    # ══════════════════════════════════════════════════════════════
    # SEC 1: System Overview
    # ══════════════════════════════════════════════════════════════
    section_header("1","ภาพรวมระบบ","System Overview & Architecture","📦")
    story.append(Paragraph(
        "ระบบบริหารคลังอะไหล่ รพ.นครพิงค์ เป็น <b>Static Web Application</b> ที่โฮสต์บน "
        "GitHub Pages เชื่อมต่อฐานข้อมูล Google Sheets ผ่าน Google Apps Script (GAS) "
        "และส่งการแจ้งเตือนผ่าน LINE OA โดยข้อมูลทั้งหมดจัดเก็บอย่างปลอดภัยบนคลาวด์", sBody))
    story.append(Spacer(1,4*mm))

    arch = [["โมดูล","รายละเอียด","มาตรฐาน"],
            ["Frontend (SPA)","HTML/CSS/JavaScript — GitHub Pages","ISO 9241"],
            ["Backend API","Google Apps Script Web App","OWASP"],
            ["Database","Google Sheets (6 ชีต)","ISO 27001"],
            ["Notification","LINE OA — Messaging API","—"],
            ["Authentication","Password-based + Role RBAC","NIST 800-63"]]
    at = Table([[Paragraph(c,sTH if i==0 else sTD) for c in r]
                for i,r in enumerate(arch)], colWidths=[45*mm, 95*mm, 30*mm])
    at.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(-1,0),C_BLUE),("TEXTCOLOR",(0,0),(-1,0),C_WHITE),
        ("GRID",(0,0),(-1,-1),0.5,C_BORDER),
        ("ROWBACKGROUNDS",(0,1),(-1,-1),[C_WHITE,C_LGRAY]),
        ("TOPPADDING",(0,0),(-1,-1),5),("BOTTOMPADDING",(0,0),(-1,-1),5),
    ]))
    story.append(at)
    story.append(Spacer(1,4*mm))
    story.append(Paragraph(bold("โครงสร้างฐานข้อมูล Google Sheets (6 ชีต)"), sH3))
    sheets = [["ชีต","หน้าที่","แถวหัว"],
              ["Parts","รายการอะไหล่ทั้งหมด","2 แถว (EN + TH)"],
              ["Transactions","รายการรับ/เบิก/ยืม/คืน","2 แถว (EN + TH)"],
              ["Users","บัญชีผู้ใช้และสิทธิ์","2 แถว (EN + TH)"],
              ["PM_Plans","แผนบำรุงรักษา","2 แถว (EN + TH)"],
              ["LINE_Logs","บันทึกการส่งข้อความ LINE","2 แถว (EN + TH)"],
              ["LINE_IDs","LINE User/Group IDs","2 แถว (EN + TH)"]]
    st2 = Table([[Paragraph(c,sTH if i==0 else sTD) for c in r]
                 for i,r in enumerate(sheets)], colWidths=[40*mm, 100*mm, 30*mm])
    st2.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(-1,0),C_NAVY),("TEXTCOLOR",(0,0),(-1,0),C_WHITE),
        ("GRID",(0,0),(-1,-1),0.5,C_BORDER),
        ("ROWBACKGROUNDS",(0,1),(-1,-1),[C_WHITE,C_LGRAY]),
        ("TOPPADDING",(0,0),(-1,-1),5),("BOTTOMPADDING",(0,0),(-1,-1),5),
    ]))
    story.append(st2)
    story.append(PageBreak())

    # ══════════════════════════════════════════════════════════════
    # SEC 2: Login
    # ══════════════════════════════════════════════════════════════
    section_header("2","การเข้าสู่ระบบ","Login & Authentication","🔐")
    story.append(Paragraph(
        "ระบบใช้การยืนยันตัวตนด้วย <b>รหัสผ่าน (Password)</b> พร้อมระบบบริหารสิทธิ์ "
        "แบบ RBAC (Role-Based Access Control) 3 ระดับ", sBody))
    story.append(Spacer(1,3*mm))
    story += img_flow("login", "หน้าจอเข้าสู่ระบบ แสดงช่องกรอกรหัสผ่านและสิทธิ์แต่ละบทบาท", 120)
    story += img_flow("genuser1", "หน้าเข้าสู่ระบบ มุมมอง General User", 100)
    story.append(Paragraph(bold("ระดับสิทธิ์การใช้งาน"), sH3))
    roles = [["บทบาท","จำนวนสูงสุด","รหัสผ่านเริ่มต้น","ขอบเขตการใช้งาน"],
             ["Admin","3 บัญชี","(กำหนดโดยผู้ดูแล)","เข้าถึงทุกฟังก์ชัน รวม Settings"],
             ["Stock Manager","3 บัญชี","(กำหนดโดย Admin)","จัดการคลัง PM แผนงบประมาณ"],
             ["General User","10 บัญชี","(กำหนดโดย Admin)","บันทึกรายการ ดูประวัติ"]]
    rt = Table([[Paragraph(c,sTH if i==0 else sTD) for c in r]
                for i,r in enumerate(roles)], colWidths=[35*mm, 25*mm, 40*mm, 60*mm])
    rt.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(-1,0),C_BLUE),("TEXTCOLOR",(0,0),(-1,0),C_WHITE),
        ("GRID",(0,0),(-1,-1),0.5,C_BORDER),
        ("ROWBACKGROUNDS",(0,1),(-1,-1),[C_WHITE,C_LGRAY]),
        ("TOPPADDING",(0,0),(-1,-1),5),("BOTTOMPADDING",(0,0),(-1,-1),5),
    ]))
    story.append(rt)
    story.append(Spacer(1,4*mm))
    story.append(step_table([
        ("เปิด Browser","ไปที่ URL: https://surasak4974buem.github.io/V3/"),
        ("กรอกรหัสผ่าน","ใส่รหัสผ่านที่ได้รับจาก Admin ในช่อง 'รหัสผ่าน'"),
        ("กด 'เข้าสู่ระบบ'","ระบบตรวจสอบสิทธิ์และนำทางสู่แดชบอร์ด"),
        ("ระบบแสดงชื่อและสิทธิ์","มุมซ้ายล่างแสดง ชื่อผู้ใช้ (บทบาท)"),
    ]))
    story.append(Spacer(1,3*mm))
    story.append(info_box("⚠️ ข้อควรระวัง",
        "ห้ามแชร์รหัสผ่านกับผู้อื่น — Admin สามารถตั้ง Reset รหัสผ่านได้จากเมนู 'จัดการสิทธิ์ผู้ใช้' "
        "กรณีลืมรหัสผ่านให้ติดต่อ Admin โดยตรง", C_RED))
    story.append(PageBreak())

    # ══════════════════════════════════════════════════════════════
    # SEC 3: Dashboard
    # ══════════════════════════════════════════════════════════════
    section_header("3","แดชบอร์ดภาพรวม","Dashboard Overview","📊")
    story.append(Paragraph(
        "แดชบอร์ดแสดงข้อมูลสรุปสถานะคลังอะไหล่แบบ Real-time ประกอบด้วย "
        "KPI Cards, กราฟวิเคราะห์, และสถานะงบประมาณประจำปี", sBody))
    story += img_flow("dashboard", "หน้าแดชบอร์ดแสดง KPI Cards, กราฟ ABC Analysis และสถานะงบประมาณ", 160)
    story.append(Paragraph(bold("องค์ประกอบหลักของแดชบอร์ด"), sH3))
    dash_items = [
        ("📦 รายการอะไหล่ทั้งหมด","แสดงจำนวน SKU อะไหล่ที่ลงทะเบียนในระบบ"),
        ("⚠️ รายการที่ต้องสั่งซื้อ","อะไหล่ที่สต็อกต่ำกว่าจุดสั่งซื้อ (Reorder Point)"),
        ("📅 อะไหล่ใกล้หมดอายุ","แสดงรายการที่จะหมดอายุภายใน 90 วัน"),
        ("💳 ยอดใช้จ่ายเดือนนี้","มูลค่าอะไหล่ที่เบิกออกในเดือนปัจจุบัน"),
        ("📈 กราฟ ABC Analysis","สัดส่วน A:B:C ของมูลค่าการใช้อะไหล่"),
        ("💰 สถานะงบประมาณ","แถบความคืบหน้าการใช้งบประมาณแต่ละโครงการ"),
    ]
    for icon_title, desc in dash_items:
        story.append(Paragraph(f"• {bold(icon_title)}  — {desc}", sBullet))
    story.append(PageBreak())

    # ══════════════════════════════════════════════════════════════
    # SEC 4: Parts Catalog
    # ══════════════════════════════════════════════════════════════
    section_header("4","คลังอะไหล่และอุปกรณ์","Spare Parts Catalog","🗄️")
    story.append(Paragraph(
        "โมดูลจัดการรายชื่ออะไหล่ สเปกทางเทคนิค จุดสั่งซื้อ และระดับ ABC Classification "
        "ผู้ใช้สิทธิ์ Admin หรือ Stock Manager สามารถเพิ่ม/แก้ไข/ลบรายการได้", sBody))
    story += img_flow("catalog", "หน้าคลังอะไหล่ แสดงรายการ ค้นหา กรองตามประเภท ABC", 160)
    story.append(Paragraph(bold("ข้อมูลแต่ละรายการอะไหล่"), sH3))
    fields = [["ฟิลด์","ความหมาย","ตัวอย่าง"],
              ["รหัสอะไหล่","รหัสเฉพาะในระบบ","SR-SPO2-001"],
              ["ชื่ออะไหล่","ชื่อภาษาไทย/อังกฤษ","สายวัด SpO2 แบบ Reusable"],
              ["ยี่ห้อ/รุ่น","Manufacturer & Model","Mindray / SpO2 Cable Type A"],
              ["ราคาต่อหน่วย","ราคาซื้อเฉลี่ย (บาท)","1,500.00"],
              ["สต็อกปัจจุบัน","จำนวนคงเหลือในคลัง","25 ชิ้น"],
              ["จุดสั่งซื้อ (ROP)","ปริมาณต่ำสุดก่อนสั่ง","5 ชิ้น"],
              ["ระดับ ABC","A=สูง, B=กลาง, C=ต่ำ (มูลค่า)","A"],
              ["วันหมดอายุ","วันที่ lot นี้หมดอายุ (พ.ศ.)","15/6/2570"],
              ["ซัพพลายเออร์","ชื่อร้าน/บริษัท","บริษัท ไทยเมดิเทค จำกัด"]]
    ft = Table([[Paragraph(c,sTH if i==0 else sTD) for c in r]
                for i,r in enumerate(fields)], colWidths=[40*mm, 80*mm, 50*mm])
    ft.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(-1,0),C_BLUE),("TEXTCOLOR",(0,0),(-1,0),C_WHITE),
        ("GRID",(0,0),(-1,-1),0.5,C_BORDER),
        ("ROWBACKGROUNDS",(0,1),(-1,-1),[C_WHITE,C_LGRAY]),
        ("TOPPADDING",(0,0),(-1,-1),5),("BOTTOMPADDING",(0,0),(-1,-1),5),
    ]))
    story.append(ft)
    story.append(Spacer(1,4*mm))
    story.append(step_table([
        ("คลิก 'เพิ่มอะไหล่ใหม่'","ปุ่มสีน้ำเงินด้านบนขวา (Admin/Stock Manager เท่านั้น)"),
        ("กรอกข้อมูล","กรอกรหัส ชื่อ ราคา จุดสั่งซื้อ วันหมดอายุ ซัพพลายเออร์"),
        ("กด 'บันทึก'","ระบบแสดง Dialog ยืนยัน — กด 'ยืนยัน' เพื่อบันทึก"),
        ("ระบบคำนวณ ABC อัตโนมัติ","จัดกลุ่ม A/B/C จากมูลค่าการใช้งาน"),
    ]))
    story.append(Spacer(1,3*mm))
    story.append(info_box("💡 ABC Analysis",
        "ระบบคำนวณ ABC โดยอัตโนมัติ: A = 20% รายการ ที่คิดเป็น 80% มูลค่า, "
        "B = 30% รายการ, C = 50% รายการที่เหลือ ตาม Pareto Principle", C_BLUE))
    story.append(PageBreak())

    # ══════════════════════════════════════════════════════════════
    # SEC 5: Transactions
    # ══════════════════════════════════════════════════════════════
    section_header("5","บันทึกรับเข้า / เบิกออก","Transactions","📝")
    story.append(Paragraph(
        "โมดูลบันทึกการเคลื่อนไหวของอะไหล่ ได้แก่ รับเข้า (Receive), เบิกออก (Issue), "
        "ยืม (Borrow), คืน (Return), ปรับสต็อก (Adjust), โอนย้าย (Transfer)", sBody))
    story += img_flow("transaction", "ฟอร์มบันทึกรายการรับเข้า/เบิกออก พร้อมช่องเลขซีเรียลเครื่องและเลขซ่อม", 160)
    story.append(Paragraph(bold("ประเภทรายการที่รองรับ"), sH3))
    tx_types = [["ประเภท","ความหมาย","ผลต่อสต็อก"],
                ["Receive","รับอะไหล่เข้าคลังจากการจัดซื้อ","+"],
                ["Issue","เบิกอะไหล่ออกเพื่อซ่อมบำรุง","–"],
                ["Borrow","ยืมอะไหล่ชั่วคราว (ยังไม่ตัดสต็อก)","– (ชั่วคราว)"],
                ["Return","คืนอะไหล่ที่ยืม","+"],
                ["Adjust","ปรับยอดสต็อกจากการตรวจนับ","±"],
                ["Transfer","โอนย้ายระหว่างแผนก/ห้อง","– ต้นทาง / + ปลายทาง"]]
    txt = Table([[Paragraph(c,sTH if i==0 else sTD) for c in r]
                 for i,r in enumerate(tx_types)], colWidths=[30*mm,110*mm,25*mm])
    txt.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(-1,0),C_NAVY),("TEXTCOLOR",(0,0),(-1,0),C_WHITE),
        ("GRID",(0,0),(-1,-1),0.5,C_BORDER),
        ("ROWBACKGROUNDS",(0,1),(-1,-1),[C_WHITE,C_LGRAY]),
        ("TOPPADDING",(0,0),(-1,-1),5),("BOTTOMPADDING",(0,0),(-1,-1),5),
    ]))
    story.append(txt)
    story.append(Spacer(1,4*mm))
    story.append(Paragraph(bold("ข้อมูลบังคับกรอก (Validation)"), sH3))
    for item in ["เลขซีเรียลเครื่อง (Equipment Serial No.) — ต้องกรอกทุกครั้ง",
                 "เลขที่ใบซ่อม (Repair Order No.) — ต้องกรอกทุกครั้ง",
                 "แผนก/ห้อง — ระบุต้นทางหรือปลายทางของอะไหล่",
                 "ผู้ดำเนินการ — ระบบดึงชื่อจากบัญชีที่ login อัตโนมัติ"]:
        story.append(Paragraph(f"• {item}", sBullet))
    story.append(Spacer(1,3*mm))
    story.append(info_box("🔔 LINE Notification",
        "เมื่อบันทึกรายการ 'Issue' ที่ทำให้สต็อกต่ำกว่า Reorder Point "
        "ระบบจะส่งแจ้งเตือนทาง LINE OA อัตโนมัติ (หากเปิดใช้งาน)", C_GREEN))
    story.append(PageBreak())

    # ══════════════════════════════════════════════════════════════
    # SEC 6: Ledger
    # ══════════════════════════════════════════════════════════════
    section_header("6","ประวัติรายการคลัง","Ledger Logs","📋")
    story.append(Paragraph(
        "บันทึกประวัติการเคลื่อนไหวอะไหล่ทั้งหมด สามารถค้นหา กรองตามวันที่ ประเภท "
        "และ Export เป็นไฟล์ CSV สำหรับรายงาน", sBody))
    story += img_flow("ledger", "ตาราง Ledger แสดงประวัติรายการพร้อมวันที่ไทย เลขซีเรียล และเลขซ่อม", 160)
    story.append(Paragraph(bold("คอลัมน์ในตาราง Ledger"), sH3))
    for col in ["วันที่-เวลา — แสดงในรูปแบบ วัน/เดือน/ปี พ.ศ. ตามมาตรฐานไทย",
                "ประเภทรายการ — Receive / Issue / Borrow / Return / Adjust",
                "รหัสอะไหล่ และชื่ออะไหล่",
                "จำนวน — ปริมาณที่รับ/เบิก/ยืม/คืน",
                "เลขซีเรียลเครื่อง — หมายเลขเครื่องที่เกี่ยวข้อง",
                "เลขที่ใบซ่อม — เลขอ้างอิงใบงานซ่อม",
                "ผู้ดำเนินการ — ชื่อเจ้าหน้าที่ที่บันทึก"]:
        story.append(Paragraph(f"• {col}", sBullet))
    story.append(Spacer(1,3*mm))
    story.append(info_box("📥 Export CSV",
        "กดปุ่ม 'Export CSV' เพื่อดาวน์โหลดประวัติรายการทั้งหมดในรูปแบบ "
        "ไฟล์ .csv ที่เปิดได้ด้วย Microsoft Excel และ Google Sheets", C_TEAL))
    story.append(PageBreak())

    # ══════════════════════════════════════════════════════════════
    # SEC 7: PM Planning
    # ══════════════════════════════════════════════════════════════
    section_header("7","แผนบำรุงรักษาเชิงรุก","Preventive Maintenance Planning","🔧")
    story.append(Paragraph(
        "โมดูลวางแผนการบำรุงรักษาเชิงป้องกัน (Preventive Maintenance) "
        "พร้อมเครื่องมือคาดการณ์ความต้องการอะไหล่ล่วงหน้า ตาม ISO 13485:2016 ข้อ 7.5", sBody))
    story += img_flow("planning", "ตาราง PM Schedule แสดงแผนซ่อมบำรุงรายเครื่อง ปุ่มเพิ่ม/แก้ไข/ลบ", 160)
    story.append(Paragraph(bold("ข้อมูลในแผน PM แต่ละรายการ"), sH3))
    pm_fields = [["ฟิลด์","รายละเอียด"],
                 ["ชื่อเครื่อง","ชื่อเครื่องมือแพทย์ที่จะบำรุงรักษา"],
                 ["รหัสเครื่อง (Serial)","หมายเลขซีเรียลเครื่อง"],
                 ["แผนก/ห้อง","แผนกที่เครื่องอยู่"],
                 ["วันนัด PM","วันที่กำหนดทำ PM (ปฏิทินไทย)"],
                 ["สถานะ","Pending / In Progress / Completed"],
                 ["อะไหล่ที่ต้องใช้","รายการและจำนวนอะไหล่ที่ต้องเปลี่ยน"],
                 ["หมายเหตุ","รายละเอียดเพิ่มเติม"]]
    pmt = Table([[Paragraph(c,sTH if i==0 else sTD) for c in r]
                 for i,r in enumerate(pm_fields)], colWidths=[50*mm, 115*mm])
    pmt.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(-1,0),C_BLUE),("TEXTCOLOR",(0,0),(-1,0),C_WHITE),
        ("GRID",(0,0),(-1,-1),0.5,C_BORDER),
        ("ROWBACKGROUNDS",(0,1),(-1,-1),[C_WHITE,C_LGRAY]),
        ("TOPPADDING",(0,0),(-1,-1),5),("BOTTOMPADDING",(0,0),(-1,-1),5),
    ]))
    story.append(pmt)
    story.append(Spacer(1,4*mm))
    story.append(step_table([
        ("กด 'เพิ่มแผน PM'","ปุ่มสีน้ำเงินด้านบน (Admin/Stock Manager)"),
        ("กรอกข้อมูลเครื่อง","ชื่อเครื่อง Serial แผนก วันนัด"),
        ("เพิ่มรายการอะไหล่","กด '+เพิ่มอะไหล่' และระบุรหัส จำนวน"),
        ("บันทึก","ระบบแสดง Dialog ยืนยัน — กด 'ยืนยัน'"),
        ("ติดตามสถานะ","แก้ไขสถานะ Pending → In Progress → Completed"),
    ]))
    story.append(PageBreak())

    # ══════════════════════════════════════════════════════════════
    # SEC 8: Procurement
    # ══════════════════════════════════════════════════════════════
    section_header("8","ใบขอจัดซื้ออะไหล่","Procurement & Requisition","🛒")
    story.append(Paragraph(
        "โมดูลติดตามอะไหล่ที่สต็อกต่ำกว่าจุดสั่งซื้อ และสร้างใบขอเสนอซื้อ (บันทึกข้อความ) "
        "อัตโนมัติตามระเบียบการจัดซื้อจัดจ้างภาครัฐ", sBody))
    story += img_flow("procurement", "หน้าใบขอจัดซื้อ แสดงรายการอะไหล่ต่ำกว่า ROP และบันทึกข้อความขอซื้อ", 160)
    story.append(Paragraph(bold("การทำงานของโมดูลใบขอจัดซื้อ"), sH3))
    for item in [
        "ระบบสแกนสต็อกทุกรายการเทียบกับ Reorder Point (ROP) โดยอัตโนมัติ",
        "รายการที่ต่ำกว่า ROP จะแสดงในตาราง 'อะไหล่ที่ต้องจัดซื้อ' พร้อมสีแดง",
        "กด 'สร้างใบขอซื้อ' เพื่อสร้างบันทึกข้อความ PDF รูปแบบราชการ",
        "บันทึกข้อความมีหัวข้อ วัตถุประสงค์ รายการ ราคา และลายมือชื่อผู้อนุมัติ",
        "บันทึกข้อมูลการจัดซื้อสำหรับ Audit Trail"]:
        story.append(Paragraph(f"• {item}", sBullet))
    story.append(PageBreak())

    # ══════════════════════════════════════════════════════════════
    # SEC 9: Budget Management
    # ══════════════════════════════════════════════════════════════
    section_header("9","แผนงบประมาณประจำปี","Annual Budget Management","💰")
    story.append(Paragraph(
        "โมดูลบริหารงบประมาณจัดซื้ออะไหล่ประจำปีงบประมาณ (ตุลาคม — กันยายน) "
        "คำนวณยอดที่ใช้ไปแล้ว คงเหลือ และความคืบหน้าแต่ละโครงการแบบ Real-time", sBody))
    story += img_flow("budget", "หน้าแผนงบประมาณ แสดง Summary Cards และตาราง CRUD รายการงบ", 160)
    story.append(Paragraph(bold("ส่วนประกอบของหน้างบประมาณ"), sH3))
    for item in [
        "Summary Cards 4 ใบ: งบทั้งหมด / ใช้ไปแล้ว / คงเหลือ / % ที่ใช้ไป",
        "ตาราง CRUD: เพิ่ม/แก้ไข/ลบรายการงบประมาณแต่ละโครงการ",
        "Progress Bar ต่อรายการ: สีเขียว (<70%) → เหลือง (70–85%) → แดง (>85%)",
        "คำนวณ 'ใช้ไปแล้ว' จากรายการ Issue ในปีงบประมาณปัจจุบันโดยอัตโนมัติ",
    ]:
        story.append(Paragraph(f"• {item}", sBullet))
    story.append(Spacer(1,4*mm))
    story.append(step_table([
        ("เปิดเมนู 'แผนงบประมาณประจำปี'","คลิกเมนูด้านซ้าย (Admin/Stock Manager)"),
        ("ดู Summary Cards","ตรวจสอบยอดรวมงบประมาณ ใช้ไปแล้ว และคงเหลือ"),
        ("เพิ่มรายการงบ","กด '+ เพิ่มรายการงบ' — กรอกชื่อโครงการ ประเภท วงเงิน"),
        ("แก้ไขรายการ","คลิก 'แก้ไข' — ระบบแสดง Dialog ยืนยัน"),
        ("ลบรายการ","คลิก 'ลบ' — ระบบแสดง Dialog ยืนยันสีแดง"),
    ]))
    story.append(Spacer(1,3*mm))
    story.append(info_box("📅 ปีงบประมาณไทย",
        "ระบบใช้ปีงบประมาณไทย (1 ตุลาคม — 30 กันยายน) โดยอัตโนมัติ "
        "ตัวอย่าง: งบประมาณปี 2568 = 1 ต.ค. 2567 ถึง 30 ก.ย. 2568", C_GOLD))
    story.append(PageBreak())

    # ══════════════════════════════════════════════════════════════
    # SEC 10: User Management
    # ══════════════════════════════════════════════════════════════
    section_header("10","จัดการสิทธิ์ผู้ใช้","User Management","👥")
    story.append(Paragraph(
        "เฉพาะ Admin เท่านั้นที่เข้าถึงหน้านี้ได้ — สามารถสร้างบัญชี กำหนดสิทธิ์ "
        "ตั้งรหัสผ่าน ระงับ/เปิดใช้งานบัญชี และดูสถิติการ Login", sBody))
    story += img_flow("usermgmt", "หน้าจัดการสิทธิ์ผู้ใช้ แสดงรายชื่อบัญชีและปุ่มจัดการ", 160)
    story += img_flow("genuser2", "มุมมอง General User: เห็นเฉพาะเมนูที่มีสิทธิ์", 150)
    story.append(Paragraph(bold("ขีดจำกัดจำนวนบัญชี"), sH3))
    for item in ["Admin: สูงสุด 3 บัญชี",
                 "Stock Manager: สูงสุด 3 บัญชี",
                 "General User: สูงสุด 10 บัญชี",
                 "รวมทั้งระบบ: สูงสุด 16 บัญชีพร้อมกัน"]:
        story.append(Paragraph(f"• {item}", sBullet))
    story.append(Spacer(1,3*mm))
    story.append(step_table([
        ("คลิก 'สร้างบัญชีใหม่'","กรอกชื่อจริง เลือกบทบาท ตั้งรหัสผ่าน"),
        ("กด 'บันทึก'","Dialog ยืนยันปรากฏ — ยืนยันเพื่อสร้างบัญชี"),
        ("แจ้งรหัสผ่านแก่ผู้ใช้","แจ้งผ่านช่องทางที่ปลอดภัย (ห้ามส่งทาง LINE ส่วนตัว)"),
        ("เปลี่ยนรหัสผ่าน","Admin คลิก 'เปลี่ยนรหัส' ต่อรายการ"),
        ("ระงับบัญชี","Toggle สถานะ Active/Suspended หากจำเป็น"),
    ]))
    story.append(PageBreak())

    # ══════════════════════════════════════════════════════════════
    # SEC 11: Settings
    # ══════════════════════════════════════════════════════════════
    section_header("11","ตั้งค่าระบบ","Cloud & LINE Settings","⚙️")
    story.append(Paragraph(
        "หน้าตั้งค่าการเชื่อมต่อฐานข้อมูล Google Sheets (ผ่าน GAS Web App URL) "
        "และระบบแจ้งเตือน LINE OA — เฉพาะ Admin เท่านั้น", sBody))
    story += img_flow("settings", "หน้าตั้งค่า Cloud & LINE แสดงช่อง GAS URL, Token, Group ID", 160)
    story += img_flow("stockmgr", "มุมมอง Stock Manager: เห็นเมนูคลัง PM และงบประมาณ", 150)
    story.append(Paragraph(bold("รายการตั้งค่าสำคัญ"), sH3))
    settings_items = [
        ("GAS Web App URL","ลิงก์ Google Apps Script ที่ Deploy แล้ว",
         "https://script.google.com/macros/s/..."),
        ("LINE Channel Access Token","Token ความยาว 172 ตัวอักษรจาก LINE Developers",
         "เก็บใน localStorage เท่านั้น — ไม่อัปโหลด GitHub"),
        ("LINE Group ID","รหัสกลุ่ม LINE OA ที่จะรับการแจ้งเตือน","Cxxxxxxxx"),
        ("เปิด/ปิด LINE","Toggle เปิด-ปิดการส่งข้อความ LINE","เปิด = ✅"),
    ]
    for name, desc, ex in settings_items:
        story.append(Paragraph(f"• {bold(name)} — {desc}", sBullet))
        story.append(Paragraph(f"  ตัวอย่าง: {ex}", sSmall))
    story.append(Spacer(1,3*mm))
    story.append(info_box("🔒 ความปลอดภัย LINE Token",
        "LINE Channel Access Token และ Group ID จัดเก็บใน Browser localStorage เท่านั้น "
        "ไม่มีการอัปโหลดขึ้น GitHub ซึ่งเป็น Public Repository "
        "ห้ามบันทึก Token ในไฟล์ใดๆ ที่ Push ขึ้น Repository", C_RED))
    story.append(Spacer(1,4*mm))
    story.append(Paragraph(bold("ขั้นตอนการตั้งค่า Cloud Sync ครั้งแรก"), sH3))
    story.append(step_table([
        ("สร้าง Google Sheet","สร้างไฟล์ใหม่บนบัญชี Google ของโรงพยาบาล"),
        ("เปิด Apps Script","เมนู ส่วนขยาย → Apps Script"),
        ("วางโค้ด code.gs","Copy โค้ดจาก GitHub repo แล้ววางใน Apps Script"),
        ("Deploy Web App","Deploy → New deployment → เลือก Anyone"),
        ("Copy URL","นำ URL ที่ได้มาวางใน GAS URL Settings"),
        ("ทดสอบ Sync","กด 'Cloud Sync' เพื่อทดสอบการเชื่อมต่อ"),
    ]))
    story.append(PageBreak())

    # ══════════════════════════════════════════════════════════════
    # SEC 12: Permissions & FAQ
    # ══════════════════════════════════════════════════════════════
    section_header("12","สิทธิ์การใช้งาน & แก้ปัญหา","Permissions & Troubleshooting","🛡️")
    story.append(Paragraph(bold("ตารางสิทธิ์การเข้าถึงแต่ละโมดูล (Access Control Matrix)"), sH3))
    story.append(perm_table())
    story.append(Spacer(1,6*mm))
    story.append(Paragraph(bold("คำถามที่พบบ่อย (FAQ)"), sH3))
    faqs = [
        ("LINE Token หายทุกครั้งที่เปิดหน้าเว็บใหม่",
         "LINE Token เก็บใน localStorage ซึ่งขึ้นอยู่กับ Browser Profile "
         "หากล้างประวัติ Browser หรือใช้ Incognito Mode ข้อมูลจะหาย "
         "→ ตั้งค่า Token ใหม่ทุกครั้งหรือใช้ Browser Profile ถาวร"),
        ("ข้อมูลหายหลังรีเฟรชหน้า",
         "ข้อมูลบางส่วนเก็บใน sessionStorage ซึ่งล้างอัตโนมัติเมื่อปิด Tab "
         "→ กด 'Cloud Sync' ก่อนบันทึกงาน เพื่อ Pull ข้อมูลล่าสุดจาก Sheets"),
        ("ส่ง LINE ไม่ได้",
         "ตรวจสอบ: 1) Token ครบ 172 ตัวอักษร  2) Group ID ถูกต้อง  "
         "3) บอท LINE ถูกเพิ่มเข้ากลุ่มแล้ว  4) Token ยังไม่หมดอายุ"),
        ("วันหมดอายุแสดงเป็น UTC หรือรูปแบบผิด",
         "ตรวจสอบว่าค่าวันที่ใน Google Sheets เป็นรูปแบบ YYYY-MM-DD "
         "ระบบจะแปลงเป็นปฏิทินไทยอัตโนมัติ (พ.ศ. = ค.ศ.+543)"),
        ("Cloud Sync ล้มเหลว (Error 500)",
         "ตรวจสอบ: 1) GAS URL ถูกต้อง  2) Deploy เป็น 'Anyone'  "
         "3) Google Sheets ไม่ได้ถูก Restrict permissions  4) ลอง Deploy ใหม่"),
    ]
    for q, a in faqs:
        story.append(KeepTogether([
            Paragraph(f"❓ {bold(q)}", sH3),
            Paragraph(f"→ {a}", sBody),
            Spacer(1,4*mm),
        ]))
    story.append(Spacer(1,4*mm))
    story.append(Paragraph(bold("การติดต่อสอบถามและรายงานปัญหา"), sH3))
    story.append(Paragraph(
        "• GitHub Issues: https://github.com/surasak4974buem/V3/issues", sBullet))
    story.append(Paragraph(
        "• ติดต่อกลุ่มงานวิศวกรรมการแพทย์ โรงพยาบาลนครพิงค์ จ.เชียงใหม่", sBullet))

    # ── Back cover ──────────────────────────────────────────────
    next_tpl("Back")
    story.append(Paragraph(" ", S("sp2")))

    doc.build(story)
    print(f"[PDF] OK  {out}")
    return out


# ════════════════════════════════════════════════════════════════
#  PART 2 – DOCX  (python-docx)
# ════════════════════════════════════════════════════════════════
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

def hex_color(h):
    h = h.lstrip("#")
    return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

def set_cell_bg(cell, hex_str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_str.lstrip("#"))
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top","left","bottom","right"):
        if edge in kwargs:
            tag = OxmlElement(f"w:{edge}")
            tag.set(qn("w:val"),  kwargs[edge].get("val","single"))
            tag.set(qn("w:sz"),   str(kwargs[edge].get("sz", 4)))
            tag.set(qn("w:space"),"0")
            tag.set(qn("w:color"),kwargs[edge].get("color","000000"))
            tcBorders.append(tag)
    tcPr.append(tcBorders)

def add_heading(doc, text, level=1, color="003580"):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = hex_color(color)
        run.font.name = "Cordia New"
        run.font.bold = True
    return p

def add_para(doc, text, size=14, bold_=False, italic=False,
             color="1a2332", align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.3
    run = p.add_run(text)
    run.font.name = "Cordia New"
    run.font.size = Pt(size)
    run.font.bold = bold_
    run.font.italic = italic
    run.font.color.rgb = hex_color(color)
    return p

def add_bullet(doc, text, size=13):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f"• {text}")
    run.font.name = "Cordia New"
    run.font.size = Pt(size)
    return p

def add_img(doc, key, caption="", width_cm=14.5):
    p_path = IMGS.get(key)
    if not p_path or not p_path.exists():
        return
    doc.add_picture(str(p_path), width=Cm(width_cm))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cp = doc.add_paragraph(f"ภาพที่: {caption}")
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(8)
        for r in cp.runs:
            r.font.size = Pt(11)
            r.font.italic = True
            r.font.color.rgb = hex_color("6B7A8D")
            r.font.name = "Cordia New"

def add_table_header_row(table, headers, bg="003580"):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.text = ""
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255,255,255)
        run.font.name = "Cordia New"
        run.font.size = Pt(12)

def add_data_row(table, values, shade_even=True, row_idx=1):
    row = table.add_row()
    bg = "F5F7FA" if (row_idx % 2 == 0 and shade_even) else "FFFFFF"
    for i, v in enumerate(values):
        cell = row.cells[i]
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        run = p.add_run(v)
        run.font.name = "Cordia New"
        run.font.size = Pt(12)
    return row

def add_section_divider(doc, num, th, en):
    doc.add_paragraph()
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.columns[0].width = Cm(2)
    tbl.columns[1].width = Cm(14)
    row = tbl.rows[0]
    c0, c1 = row.cells[0], row.cells[1]
    set_cell_bg(c0, "003580")
    c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p0.add_run(str(num))
    r0.font.bold = True; r0.font.size = Pt(18)
    r0.font.color.rgb = RGBColor(255,255,255); r0.font.name = "Cordia New"
    c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p1 = c1.paragraphs[0]
    r1 = p1.add_run(th)
    r1.font.bold = True; r1.font.size = Pt(18)
    r1.font.color.rgb = hex_color("003580"); r1.font.name = "Cordia New"
    p2 = c1.add_paragraph(en)
    p2.paragraph_format.space_after = Pt(0)
    for r in p2.runs:
        r.font.size = Pt(11); r.font.color.rgb = hex_color("6B7A8D")
        r.font.name = "Cordia New"
    # Gold bottom border
    from docx.oxml import OxmlElement as OE
    tbl._tbl.append(OE("w:tblPr"))
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def build_docx():
    doc = Document()
    # Page setup A4
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = section.right_margin = Cm(2.5)
    section.top_margin  = section.bottom_margin = Cm(2.0)

    # Default font
    doc.styles["Normal"].font.name = "Cordia New"
    doc.styles["Normal"].font.size = Pt(14)

    # ── COVER PAGE ───────────────────────────────────────────────
    cover_tbl = doc.add_table(rows=1, cols=1)
    cover_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cover_tbl.columns[0].width = Cm(16)
    cc = cover_tbl.rows[0].cells[0]
    set_cell_bg(cc, "003580")
    cc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    def cover_line(text, size, color="FFFFFF", bold_=False, space=4, center=True):
        p = cc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(space)
        r = p.add_run(text)
        r.font.name = "Cordia New"; r.font.size = Pt(size)
        r.font.color.rgb = hex_color(color); r.font.bold = bold_
        return p

    # Remove default empty para in cell
    cc.paragraphs[0].clear()
    cover_line("", 10, space=20)
    cover_line("✚", 48, color="FFFFFF", space=4)
    cover_line("โรงพยาบาลนครพิงค์  Nakornping Hospital", 16, color="7ECEF4", bold_=True, space=2)
    cover_line("ศูนย์เครื่องมือแพทย์  •  งานซ่อมบำรุงและคลังอะไหล่", 12, color="AED6F1", space=20)
    cover_line("คู่มือการใช้งาน", 30, color="FFFFFF", bold_=True, space=4)
    cover_line("ระบบบริหารคลังอะไหล่และอุปกรณ์", 22, color="FFFFFF", bold_=True, space=4)
    cover_line("Medical Spare Parts Management System", 14, color="90CAF9", space=20)
    cover_line(f"Version {VERSION}   |   รหัส: {DOC_CODE}", 12, color="D4E8FF", space=4)
    cover_line(f"วันที่: {TODAY_TH}   |   มาตรฐาน SMM 07-1:2024", 12, color="D4E8FF", space=30)
    cover_line("จัดทำโดย: กลุ่มงานวิศวกรรมการแพทย์  โรงพยาบาลนครพิงค์  จ.เชียงใหม่",
               11, color="90CAF9", space=10)

    doc.add_page_break()

    # ── คำนำ ────────────────────────────────────────────────────
    add_heading(doc, "คำนำ", level=1)
    add_para(doc,
        "คู่มือการใช้งาน ระบบบริหารคลังอะไหล่และอุปกรณ์ โรงพยาบาลนครพิงค์ (Version 3) "
        "จัดทำขึ้นเพื่อเป็นแนวทางสำหรับเจ้าหน้าที่ วิศวกรชีวการแพทย์ และผู้ดูแลระบบ "
        "ในการใช้งานระบบสารสนเทศเพื่อการบริหารจัดการคลังอะไหล่เครื่องมือแพทย์ "
        "ให้เป็นไปตามมาตรฐาน SMM 07-1:2024 และ ISO 13485:2016", 13, space_after=8)
    add_para(doc,
        "ระบบนี้พัฒนาในรูปแบบ Web Application ไม่ต้องติดตั้งโปรแกรม สามารถเข้าใช้ผ่าน Browser "
        "ทุกอุปกรณ์ เชื่อมต่อ Google Sheets เป็นฐานข้อมูลคลาวด์ และแจ้งเตือนผ่าน LINE OA", 13)

    add_heading(doc, "วัตถุประสงค์ของคู่มือ", level=2)
    for b in ["อธิบายขั้นตอนการใช้งานระบบในแต่ละโมดูลพร้อมภาพประกอบ",
              "กำหนดสิทธิ์การเข้าถึงและความรับผิดชอบของผู้ใช้แต่ละระดับ",
              "แนะนำการแก้ไขปัญหาเบื้องต้นและข้อควรระวัง",
              "ใช้เป็นเอกสารอ้างอิงสำหรับการฝึกอบรมบุคลากรใหม่"]:
        add_bullet(doc, b)

    doc.add_paragraph()
    meta_tbl = doc.add_table(rows=7, cols=2)
    meta_tbl.style = "Table Grid"
    headers = ["รายการ","รายละเอียด"]
    add_table_header_row(meta_tbl, headers)
    data = [("ชื่อระบบ","ระบบบริหารคลังอะไหล่และอุปกรณ์ รพ.นครพิงค์"),
            ("รหัสเอกสาร", DOC_CODE),("เวอร์ชัน",f"Version {VERSION}"),
            ("วันที่จัดทำ", TODAY_TH),
            ("ผู้จัดทำ","กลุ่มงานวิศวกรรมการแพทย์ โรงพยาบาลนครพิงค์"),
            ("มาตรฐาน","SMM 07-1:2024, ISO 13485:2016")]
    for i,(k,v) in enumerate(data):
        add_data_row(meta_tbl, [k,v], row_idx=i)

    doc.add_page_break()

    # ── สารบัญ ──────────────────────────────────────────────────
    add_heading(doc, "สารบัญ  Table of Contents", level=1)
    toc_data = [
        ("คำนำ","Preface","I"),
        ("สารบัญ","Table of Contents","II"),
        ("บทที่ 1","ภาพรวมระบบ — System Overview","3"),
        ("บทที่ 2","การเข้าสู่ระบบ — Login & Authentication","4"),
        ("บทที่ 3","แดชบอร์ดภาพรวม — Dashboard","5"),
        ("บทที่ 4","คลังอะไหล่และอุปกรณ์ — Parts Catalog","6"),
        ("บทที่ 5","บันทึกรับเข้า/เบิกออก — Transactions","7"),
        ("บทที่ 6","ประวัติรายการคลัง — Ledger Logs","8"),
        ("บทที่ 7","แผนบำรุงรักษาเชิงรุก — PM Planning","9"),
        ("บทที่ 8","ใบขอจัดซื้ออะไหล่ — Procurement","10"),
        ("บทที่ 9","แผนงบประมาณประจำปี — Budget Management","11"),
        ("บทที่ 10","จัดการสิทธิ์ผู้ใช้ — User Management","12"),
        ("บทที่ 11","ตั้งค่าระบบ — Cloud & LINE Settings","13"),
        ("บทที่ 12","สิทธิ์การใช้งาน & แก้ปัญหา — Permissions & FAQ","14"),
    ]
    toc_tbl = doc.add_table(rows=1, cols=3)
    toc_tbl.style = "Table Grid"
    add_table_header_row(toc_tbl, ["บทที่","หัวข้อ","หน้า"])
    for i,(num,title,pg) in enumerate(toc_data):
        r = toc_tbl.add_row()
        for j,v in enumerate([num,title,pg]):
            c = r.cells[j]
            set_cell_bg(c, "F5F7FA" if i%2==0 else "FFFFFF")
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j!=1 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(v)
            run.font.name = "Cordia New"; run.font.size = Pt(12)
            if j==1 and num in ("คำนำ","สารบัญ"):
                run.font.bold = True

    doc.add_page_break()

    # ── Sections 1–12 (condensed but complete) ──────────────────
    def sec(num, th, en, body_text, img_keys=[], bullets=[], steps=[], note=None):
        add_section_divider(doc, num, th, en)
        add_para(doc, body_text, 13)
        for k,cap in img_keys:
            add_img(doc, k, cap)
        if bullets:
            add_heading(doc, bullets[0], level=2)
            for b in bullets[1:]:
                add_bullet(doc, b)
        if steps:
            add_heading(doc, "ขั้นตอนการใช้งาน", level=2)
            for i,(st,sd) in enumerate(steps,1):
                add_para(doc, f"{i}. {st}", 13, bold_=True, space_after=1)
                add_para(doc, f"   → {sd}", 12, space_after=4)
        if note:
            add_para(doc, f"⚠️  {note}", 12, italic=True, color="C0392B")
        doc.add_page_break()

    sec(1,"ภาพรวมระบบ","System Overview",
        "ระบบบริหารคลังอะไหล่ รพ.นครพิงค์ เป็น Static Web Application บน GitHub Pages "
        "เชื่อมต่อ Google Sheets ผ่าน GAS Web App และแจ้งเตือนผ่าน LINE OA "
        "รองรับ 3 ระดับสิทธิ์ RBAC โดยข้อมูลจัดเก็บใน 6 Google Sheets",
        bullets=["โมดูลหลักของระบบ",
                 "บริหารคลังอะไหล่ — เพิ่ม/แก้ไข/ลบ ABC Analysis",
                 "บันทึกรายการ — รับ เบิก ยืม คืน ปรับ โอน",
                 "แผน PM — ตาราง Preventive Maintenance ประจำปี",
                 "ใบขอจัดซื้อ — สร้างเอกสารขอซื้ออัตโนมัติ",
                 "งบประมาณ — ติดตามวงเงินรายโครงการ Real-time",
                 "Cloud Sync — สำรองข้อมูลลง Google Sheets"])

    sec(2,"การเข้าสู่ระบบ","Login & Authentication",
        "ยืนยันตัวตนด้วยรหัสผ่าน (Password) พร้อม RBAC 3 ระดับ: Admin, Stock Manager, General User",
        img_keys=[("login","หน้าเข้าสู่ระบบ"),("genuser1","มุมมอง General User")],
        steps=[("เปิด URL","https://surasak4974buem.github.io/V3/"),
               ("กรอกรหัสผ่าน","รหัสที่ได้รับจาก Admin"),
               ("กด 'เข้าสู่ระบบ'","ระบบนำไปสู่แดชบอร์ดตามสิทธิ์"),
               ("ออกจากระบบ","กดปุ่ม 'ออกจากระบบ' มุมซ้ายล่าง")],
        note="ห้ามแชร์รหัสผ่าน — ติดต่อ Admin เพื่อ Reset รหัสผ่าน")

    sec(3,"แดชบอร์ดภาพรวม","Dashboard",
        "หน้าหลักแสดง KPI Cards, กราฟ ABC Analysis, กราฟแนวโน้ม และสถานะงบประมาณ",
        img_keys=[("dashboard","แดชบอร์ดภาพรวม")],
        bullets=["ส่วนประกอบของแดชบอร์ด",
                 "📦 รายการอะไหล่ทั้งหมด — จำนวน SKU",
                 "⚠️ รายการต้องสั่งซื้อ — ต่ำกว่า Reorder Point",
                 "📅 ใกล้หมดอายุ — ภายใน 90 วัน",
                 "💳 ยอดใช้จ่ายเดือนนี้",
                 "📈 กราฟ ABC Analysis",
                 "💰 สถานะงบประมาณประจำปี"])

    sec(4,"คลังอะไหล่และอุปกรณ์","Parts Catalog",
        "จัดการรายชื่ออะไหล่ สเปกทางเทคนิค จุดสั่งซื้อ ABC Classification "
        "Admin และ Stock Manager มีสิทธิ์ CRUD เต็ม",
        img_keys=[("catalog","หน้าคลังอะไหล่")],
        steps=[("คลิก 'เพิ่มอะไหล่ใหม่'","Admin/Stock Manager เท่านั้น"),
               ("กรอกข้อมูล","รหัส ชื่อ ราคา ROP วันหมดอายุ ซัพพลายเออร์"),
               ("ยืนยัน","Dialog ยืนยัน → กด 'ยืนยัน'"),
               ("ระบบคำนวณ ABC","จัดกลุ่มอัตโนมัติจากมูลค่าการใช้งาน")])

    sec(5,"บันทึกรับเข้า / เบิกออก","Transactions",
        "บันทึกการเคลื่อนไหวอะไหล่ 6 ประเภท: Receive, Issue, Borrow, Return, Adjust, Transfer "
        "บังคับกรอก Serial เครื่องและเลขที่ใบซ่อมทุกรายการ",
        img_keys=[("transaction","ฟอร์มบันทึกรายการ")],
        steps=[("เลือกประเภทรายการ","Receive / Issue / Borrow / Return / Adjust / Transfer"),
               ("เลือกอะไหล่","ค้นหาด้วยรหัสหรือชื่อ"),
               ("กรอก Serial + เลขซ่อม","บังคับกรอกทุกครั้ง"),
               ("กด 'บันทึก'","ระบบยืนยันและอัปเดตสต็อก"),
               ("LINE แจ้งเตือน","หากสต็อกต่ำกว่า ROP — ส่งแจ้งเตือนอัตโนมัติ")])

    sec(6,"ประวัติรายการคลัง","Ledger Logs",
        "บันทึกประวัติรายการทั้งหมดในรูปแบบตาราง กรองตามวันที่และประเภท Export CSV",
        img_keys=[("ledger","ตาราง Ledger Logs")],
        bullets=["คอลัมน์หลักใน Ledger",
                 "วันที่-เวลา ในรูปแบบปฏิทินไทย (พ.ศ.)",
                 "ประเภทรายการ (Receive/Issue/Borrow/...)",
                 "รหัสและชื่ออะไหล่, จำนวน",
                 "เลขซีเรียลเครื่อง และเลขที่ใบซ่อม",
                 "ผู้ดำเนินการ, แผนก"])

    sec(7,"แผนบำรุงรักษาเชิงรุก","PM Planning",
        "วางแผน PM ประจำปีตาม ISO 13485:2016 พร้อมคาดการณ์อะไหล่ที่ต้องใช้",
        img_keys=[("planning","ตาราง PM Schedule")],
        steps=[("กด 'เพิ่มแผน PM'","Admin/Stock Manager"),
               ("กรอกชื่อเครื่อง Serial แผนก","ระบุวันนัด PM"),
               ("เพิ่มรายการอะไหล่","กด '+เพิ่มอะไหล่' ระบุรหัสและจำนวน"),
               ("บันทึก","Dialog ยืนยัน"),
               ("ติดตามสถานะ","Pending → In Progress → Completed")])

    sec(8,"ใบขอจัดซื้ออะไหล่","Procurement & Requisition",
        "ระบบสแกนสต็อกเทียบ ROP อัตโนมัติ สร้างใบขอเสนอซื้อรูปแบบราชการ",
        img_keys=[("procurement","หน้าใบขอจัดซื้อ")],
        bullets=["การทำงานของโมดูลจัดซื้อ",
                 "สแกนสต็อกทุกรายการเทียบ Reorder Point อัตโนมัติ",
                 "รายการต่ำกว่า ROP แสดงสีแดงในตาราง",
                 "กด 'สร้างใบขอซื้อ' เพื่อสร้างเอกสาร PDF",
                 "รูปแบบบันทึกข้อความราชการพร้อมลายมือชื่อผู้อนุมัติ"])

    sec(9,"แผนงบประมาณประจำปี","Annual Budget Management",
        "บริหารงบประมาณรายโครงการ คำนวณยอดใช้ไปแล้ว คงเหลือ และ % ความคืบหน้า Real-time",
        img_keys=[("budget","หน้าแผนงบประมาณ")],
        steps=[("เปิดเมนู 'แผนงบประมาณประจำปี'","Admin/Stock Manager เท่านั้น"),
               ("ดู Summary Cards","งบทั้งหมด / ใช้ไป / คงเหลือ / %"),
               ("เพิ่มรายการงบ","กด '+ เพิ่มรายการงบ' — กรอกโครงการ ประเภท วงเงิน"),
               ("แก้ไข/ลบ","ปุ่มแต่ละแถว — Dialog ยืนยันทุกครั้ง")])

    sec(10,"จัดการสิทธิ์ผู้ใช้","User Management",
        "Admin จัดการบัญชีผู้ใช้ทั้งหมด กำหนดสิทธิ์ ตั้งรหัสผ่าน ระงับ/เปิดใช้งาน",
        img_keys=[("usermgmt","หน้าจัดการผู้ใช้"),("genuser2","มุมมอง General User")],
        steps=[("คลิก 'สร้างบัญชีใหม่'","กรอกชื่อ เลือกบทบาท ตั้งรหัสผ่าน"),
               ("บันทึก","Dialog ยืนยัน"),
               ("แจ้งรหัสผ่าน","ผ่านช่องทางที่ปลอดภัย"),
               ("ระงับบัญชี","Toggle Active/Suspended")])

    sec(11,"ตั้งค่าระบบ","Cloud & LINE Settings",
        "Admin ตั้งค่า GAS URL, LINE Token, Group ID — เก็บใน localStorage ไม่อัปโหลด GitHub",
        img_keys=[("settings","หน้าตั้งค่า"),("stockmgr","มุมมอง Stock Manager")],
        steps=[("กรอก GAS URL","URL จาก Google Apps Script"),
               ("กรอก LINE Token","Token 172 ตัวจาก LINE Developers"),
               ("กรอก LINE Group ID","ID กลุ่มที่รับแจ้งเตือน"),
               ("กด 'บันทึก'","ระบบเก็บใน localStorage"),
               ("ทดสอบ LINE","กด 'ทดสอบส่งข้อความ'")],
        note="LINE Token เป็นความลับ — ห้ามบันทึกในไฟล์ที่อัปโหลด GitHub")

    # ── Sec 12 ──────────────────────────────────────────────────
    add_section_divider(doc, 12, "สิทธิ์การใช้งาน & แก้ปัญหา", "Permissions & FAQ")
    add_heading(doc, "ตารางสิทธิ์การเข้าถึง (Access Control Matrix)", level=2)
    perm_headers = ["ฟังก์ชัน","Admin","Stock Manager","General User"]
    perm_rows = [
        ("แดชบอร์ด","✅","✅","✅"),
        ("คลังอะไหล่ (CRUD)","✅","✅","❌"),
        ("บันทึกรายการ","✅","✅","✅"),
        ("แผน PM","✅","✅","👁️"),
        ("ใบขอจัดซื้อ","✅","✅","❌"),
        ("แผนงบประมาณ","✅","✅","❌"),
        ("จัดการผู้ใช้","✅","❌","❌"),
        ("ตั้งค่า & Cloud","✅","❌","❌"),
    ]
    pt = doc.add_table(rows=1, cols=4)
    pt.style = "Table Grid"
    add_table_header_row(pt, perm_headers)
    for i,r in enumerate(perm_rows):
        add_data_row(pt, list(r), row_idx=i)

    add_heading(doc, "คำถามที่พบบ่อย (FAQ)", level=2)
    faqs2 = [
        ("LINE Token หายหลังรีเฟรช",
         "LINE Token เก็บใน localStorage — หากล้าง Browser หรือใช้ Incognito ต้องตั้งค่าใหม่"),
        ("ข้อมูลหายหลังปิด Tab",
         "กด 'Cloud Sync' ก่อนปิด Tab เพื่อบันทึกลง Google Sheets"),
        ("ส่ง LINE ไม่ได้",
         "ตรวจ Token ครบ 172 ตัว, Group ID ถูก, บอทอยู่ในกลุ่ม, Token ยังไม่หมดอายุ"),
        ("วันหมดอายุแสดงผิด",
         "ตรวจรูปแบบวันที่ใน Sheets: ต้องเป็น YYYY-MM-DD"),
        ("Cloud Sync Error 500",
         "ตรวจ GAS URL, Deploy เป็น Anyone, Sheets ไม่ได้ Restrict, ลอง Re-deploy"),
    ]
    for q,a in faqs2:
        add_para(doc, f"❓  {q}", 13, bold_=True, space_after=2)
        add_para(doc, f"→  {a}", 12, space_after=8)

    doc.add_page_break()

    # ── BACK COVER ───────────────────────────────────────────────
    back_tbl = doc.add_table(rows=1, cols=1)
    back_tbl.columns[0].width = Cm(16)
    bc = back_tbl.rows[0].cells[0]
    set_cell_bg(bc, "001F4D")
    bc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    def bc_line(text, size, color="FFFFFF", bold_=False, space=4):
        p = bc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(space)
        r = p.add_run(text)
        r.font.name="Cordia New"; r.font.size=Pt(size)
        r.font.color.rgb=hex_color(color); r.font.bold=bold_

    bc.paragraphs[0].clear()
    bc_line("",10,space=30)
    bc_line("ระบบบริหารคลังอะไหล่และอุปกรณ์",22,"FFFFFF",True)
    bc_line("Medical Spare Parts Management System",14,"90CAF9",space=4)
    bc_line("──────────────────────────────",11,"C8A415",space=4)
    bc_line("โรงพยาบาลนครพิงค์  จังหวัดเชียงใหม่",14,"AED6F1")
    bc_line("ศูนย์เครื่องมือแพทย์  กลุ่มงานวิศวกรรมการแพทย์",12,"AED6F1")
    bc_line("",10,space=10)
    bc_line(f"รหัสเอกสาร: {DOC_CODE}",12,"D4E8FF")
    bc_line(f"Version {VERSION}  |  {TODAY_TH}",12,"D4E8FF")
    bc_line("",10,space=10)
    bc_line("GitHub: github.com/surasak4974buem/V3",12,"7ECEF4")
    bc_line("",10,space=30)

    out = OUT_DIR / "NKP_Manual_V3.docx"
    doc.save(str(out))
    print(f"[DOCX] OK  {out}")
    return out


# ════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    print("Building PDF …")
    pdf_out  = build_pdf()
    print("Building DOCX …")
    docx_out = build_docx()
    print("\nDone!")
    print(f"  PDF  : {pdf_out}")
    print(f"  DOCX : {docx_out}")
